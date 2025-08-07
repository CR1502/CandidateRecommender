"""
File processing utilities for extracting text from various file formats.
"""

import io
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import PyPDF2
import docx
from loguru import logger
import chardet


class FileProcessor:
    """Process and extract text from various file formats."""

    def __init__(self, max_file_size_mb: int = 10):
        """
        Initialize the file processor.

        Args:
            max_file_size_mb: Maximum file size in MB
        """
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.supported_formats = ['pdf', 'docx', 'txt']

    def process_file(self, file_obj: Any, filename: str) -> Tuple[str, str]:
        """
        Process uploaded file and extract text.

        Args:
            file_obj: File object from Streamlit uploader
            filename: Name of the file

        Returns:
            Tuple of (extracted_text, candidate_name)

        Raises:
            ValueError: If file type is not supported or file is too large
        """
        try:
            # Check file size
            file_obj.seek(0, 2)  # Move to end
            file_size = file_obj.tell()
            file_obj.seek(0)  # Reset to beginning

            if file_size > self.max_file_size_bytes:
                raise ValueError(f"File size exceeds {self.max_file_size_bytes / (1024 * 1024)}MB limit")

            # Get file extension
            file_ext = filename.rsplit('.', 1)[-1].lower()

            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")

            # Extract text based on file type
            if file_ext == 'pdf':
                text = self._extract_from_pdf(file_obj)
            elif file_ext == 'docx':
                text = self._extract_from_docx(file_obj)
            elif file_ext == 'txt':
                text = self._extract_from_txt(file_obj)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

            # Extract candidate name
            from .text_cleaner import TextCleaner
            cleaner = TextCleaner()
            candidate_name = cleaner.extract_candidate_name(text, filename)

            logger.info(f"Successfully processed file: {filename}")
            return text, candidate_name

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise

    def _extract_from_pdf(self, file_obj: Any) -> str:
        """
        Extract text from PDF file.

        Args:
            file_obj: PDF file object

        Returns:
            Extracted text
        """
        try:
            pdf_reader = PyPDF2.PdfReader(file_obj)
            text_parts = []

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_from_docx(self, file_obj: Any) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_obj: DOCX file object

        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file_obj)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_from_txt(self, file_obj: Any) -> str:
        """
        Extract text from TXT file with encoding detection.

        Args:
            file_obj: TXT file object

        Returns:
            Extracted text
        """
        try:
            # Read raw bytes
            raw_data = file_obj.read()

            # Detect encoding
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'

            # Decode text
            text = raw_data.decode(encoding, errors='ignore')
            return text

        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            # Fallback to UTF-8
            try:
                file_obj.seek(0)
                return file_obj.read().decode('utf-8', errors='ignore')
            except:
                raise ValueError(f"Failed to extract text from TXT: {str(e)}")

    def process_multiple_files(self, files: List[Any]) -> List[Dict[str, Any]]:
        """
        Process multiple files and extract text from each.

        Args:
            files: List of file objects

        Returns:
            List of dictionaries with file data
        """
        results = []

        for file_obj in files:
            try:
                text, candidate_name = self.process_file(file_obj, file_obj.name)
                results.append({
                    'filename': file_obj.name,
                    'candidate_name': candidate_name,
                    'text': text,
                    'error': None
                })
            except Exception as e:
                logger.error(f"Failed to process {file_obj.name}: {e}")
                results.append({
                    'filename': file_obj.name,
                    'candidate_name': None,
                    'text': None,
                    'error': str(e)
                })

        return results

    def validate_file(self, file_obj: Any, filename: str) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing.

        Args:
            file_obj: File object
            filename: File name

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if file_ext not in self.supported_formats:
            return False, f"Unsupported file type: {file_ext}. Supported: {', '.join(self.supported_formats)}"

        # Check file size
        file_obj.seek(0, 2)
        file_size = file_obj.tell()
        file_obj.seek(0)

        if file_size > self.max_file_size_bytes:
            return False, f"File too large: {file_size / (1024 * 1024):.1f}MB. Maximum: {self.max_file_size_bytes / (1024 * 1024)}MB"

        if file_size == 0:
            return False, "File is empty"

        return True, None